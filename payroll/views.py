import datetime
import os
import io
import docx
from django.conf import settings
from django.shortcuts import render, get_object_or_404, redirect


from django.db.models import Q
from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.contrib.auth import login
from django.contrib import messages
from django.http import JsonResponse, HttpResponseForbidden, FileResponse
from django.utils import timezone
from django.core.paginator import Paginator
from .models import (
    Role, Employee, Contract, ObligationRequest, DisbursementVoucher, 
    DVPayeeDetail, WorkflowLog, ObRStatus, DVStatus, ContractType,
    GeneralPayroll, GeneralPayrollItem, GeneralPayrollStatus, Department
)

@login_required
def dashboard(request):
    user_profile = request.user.profile
    role = user_profile.role

    # Filter transactions based on role or show all
    obrs = ObligationRequest.objects.all().order_by('-transaction_date', '-id')
    dvs = DisbursementVoucher.objects.all().order_by('-transaction_date', '-id')

    # Find tasks requiring action by this specific role
    action_obrs = []
    action_dvs = []

    if role == Role.UNIT_AO:
        action_obrs = obrs.filter(status=ObRStatus.DRAFTED)
        action_dvs = dvs.filter(status=DVStatus.DRAFTED)
    elif role == Role.UNIT_HEAD:
        action_obrs = obrs.filter(status=ObRStatus.SUBMITTED)
    elif role == Role.SBO_STAFF:
        action_obrs = obrs.filter(status__in=[ObRStatus.UNIT_HEAD_APPROVED, ObRStatus.SBO_OFFICER_REVIEWED, ObRStatus.SBO_SUPERVISOR_REVIEWED])
    elif role == Role.SAO_RECEIVING:
        action_dvs = dvs.filter(status=DVStatus.UNIT_HEAD_APPROVED)
    elif role == Role.SAO_PREAUDIT:
        action_dvs = dvs.filter(status=DVStatus.SAO_RECEIVED)
    elif role == Role.SAO_PAYROLLSUP:
        action_dvs = dvs.filter(status=DVStatus.SAO_PRE_AUDIT)
    elif role == Role.SAO_CASHMONI:
        action_dvs = dvs.filter(status=DVStatus.SAO_SUPERVISOR)
    elif role == Role.SAO_DIRECTOR:
        action_dvs = dvs.filter(status=DVStatus.CASH_MONITORED)
    elif role == Role.OVPA_VP:
        action_dvs = dvs.filter(status=DVStatus.SAO_DIRECTOR)
    elif role == Role.SCO_STAFF:
        action_dvs = dvs.filter(status=DVStatus.APPROVED)

    # Compute KPI values
    total_cos = Employee.objects.count()

    stats = {
        'total_dvs': dvs.count(),
        'pending_dvs': dvs.exclude(status=DVStatus.RADA_DRAFT).count(),
        'completed_dvs': dvs.filter(status=DVStatus.RADA_DRAFT).count(),
        'total_obrs': obrs.count(),
        'total_cos': total_cos,
    }

    # Chart Data 1: Status Distribution
    status_counts = {}
    for code, label in DVStatus.choices:
        count = dvs.filter(status=code).count()
        if count > 0:
            status_counts[label] = count

    # Chart Data 2: Top 5 Requesting Departments
    dept_volumes = {}
    for dept_code, dept_label in Department.choices:
        vol = sum(dv.total_amount for dv in dvs.filter(requesting_unit=dept_code))
        if vol > 0:
            dept_volumes[dept_label] = float(vol)
    top_departments = dict(sorted(dept_volumes.items(), key=lambda x: x[1], reverse=True)[:5])

    general_payrolls = GeneralPayroll.objects.all().order_by('-created_at')

    context = {
        'obrs': obrs,
        'dvs': dvs,
        'action_obrs': action_obrs,
        'action_dvs': action_dvs,
        'stats': stats,
        'general_payrolls': general_payrolls,
        'status_chart_labels': list(status_counts.keys()),
        'status_chart_data': list(status_counts.values()),
        'dept_chart_labels': list(top_departments.keys()),
        'dept_chart_data': list(top_departments.values()),
    }
    return render(request, 'payroll/dashboard.html', context)


@login_required
def transactions_list(request):
    search_query = request.GET.get('search', '')
    transaction_type = request.GET.get('type', 'ALL')
    status_filter = request.GET.get('status', '')
    unit_filter = request.GET.get('unit', '')

    items = []

    # Query source models
    dvs = DisbursementVoucher.objects.all().order_by('-transaction_date', '-id')
    obrs = ObligationRequest.objects.all().order_by('-transaction_date', '-id')

    if search_query:
        dvs = dvs.filter(Q(dv_number__icontains=search_query) | Q(requesting_unit__icontains=search_query))
        obrs = obrs.filter(Q(obr_number__icontains=search_query) | Q(requesting_unit__icontains=search_query))

    if unit_filter:
        dvs = dvs.filter(requesting_unit=unit_filter)
        obrs = obrs.filter(requesting_unit=unit_filter)

    if status_filter:
        dvs = dvs.filter(status=status_filter)
        obrs = obrs.filter(status=status_filter)

    # Format list depending on selected transaction type
    if transaction_type == 'DV':
        for dv in dvs:
            items.append({
                'type': 'DV',
                'ref_number': dv.dv_number,
                'requesting_unit': dv.requesting_unit,
                'amount': dv.total_amount,
                'status': dv.get_status_display(),
                'status_code': dv.status,
                'date': dv.transaction_date,
                'url': f"/dv/{dv.pk}/"
            })
    elif transaction_type == 'OBR':
        for obr in obrs:
            items.append({
                'type': 'OBR',
                'ref_number': obr.obr_number,
                'requesting_unit': obr.requesting_unit,
                'amount': obr.total_amount,
                'status': obr.get_status_display(),
                'status_code': obr.status,
                'date': obr.transaction_date,
                'url': f"/obr/{obr.pk}/"
            })
    else:
        # ALL
        for dv in dvs:
            items.append({
                'type': 'DV',
                'ref_number': dv.dv_number,
                'requesting_unit': dv.requesting_unit,
                'amount': dv.total_amount,
                'status': dv.get_status_display(),
                'status_code': dv.status,
                'date': dv.transaction_date,
                'url': f"/dv/{dv.pk}/"
            })
        for obr in obrs:
            items.append({
                'type': 'OBR',
                'ref_number': obr.obr_number,
                'requesting_unit': obr.requesting_unit,
                'amount': obr.total_amount,
                'status': obr.get_status_display(),
                'status_code': obr.status,
                'date': obr.transaction_date,
                'url': f"/obr/{obr.pk}/"
            })
        # Sort combined list by date
        items.sort(key=lambda x: (x['date'] or datetime.date.min), reverse=True)

    # Page processing (20 items per page)
    paginator = Paginator(items, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Rebuilding choices lists for filtering
    departments = Department.choices
    statuses = list(DVStatus.choices) + list(ObRStatus.choices)
    
    unique_statuses = []
    seen = set()
    for code, label in statuses:
        if code not in seen:
            seen.add(code)
            unique_statuses.append((code, label))

    context = {
        'page_obj': page_obj,
        'search_query': search_query,
        'transaction_type': transaction_type,
        'status_filter': status_filter,
        'unit_filter': unit_filter,
        'departments': departments,
        'statuses': unique_statuses,
    }
    return render(request, 'payroll/transactions_list.html', context)

def switch_role(request):
    username = request.GET.get('username')
    if username:
        try:
            user = User.objects.get(username=username)
            login(request, user)
            messages.success(request, f"Switched to user: {user.get_full_name()} ({user.profile.get_role_display()})")
        except User.DoesNotExist:
            messages.error(request, "Selected demo account does not exist.")
    return redirect('dashboard')

@login_required
def create_obr_dv(request):
    if request.method == 'POST':
        # Simulate loading the BULSA template screenshots data
        now = datetime.datetime.now()
        date_str = now.strftime("%Y-%m-%d")
        
        # Create Obligation Request
        obr_no = f"ORS1-UPSA-02-2026-06-03415"
        # Check if already exists to prevent duplicate key error in prototype
        if ObligationRequest.objects.filter(obr_number=obr_no).exists():
            # Add random suffix
            suffix = now.strftime("%H%M%S")
            obr_no = f"ORS1-UPSA-02-2026-06-03415-{suffix}"
            dv_no = f"ORS1-UPSA-2026-07-03607-{suffix}"
        else:
            dv_no = f"ORS1-UPSA-2026-07-03607"

        # Create ObR
        obr = ObligationRequest.objects.create(
            obr_number=obr_no,
            total_amount=145438.85,
            status=ObRStatus.DRAFTED,
            created_by=request.user
        )
        
        # Create Workflow Log for ObR
        WorkflowLog.objects.create(
            obr=obr,
            status_from='',
            status_to=ObRStatus.DRAFTED,
            changed_by=request.user,
            remarks='Saved ObR as draft (Initiated via BULSA Upload)'
        )

        gp_id = request.POST.get('general_payroll_id')
        gp = None
        if gp_id:
            gp = get_object_or_404(GeneralPayroll, pk=gp_id)

        # Create Disbursement Voucher linked to ObR
        dv = DisbursementVoucher.objects.create(
            dv_number=dv_no,
            obr=obr,
            total_amount=145438.85,
            status=DVStatus.DRAFTED,
            created_by=request.user,
            general_payroll=gp
        )

        # Create Workflow Log for DV
        WorkflowLog.objects.create(
            dv=dv,
            status_from='',
            status_to=DVStatus.DRAFTED,
            changed_by=request.user,
            remarks='Saved DV as draft (Initiated via BULSA Upload)'
        )

        # Create Payee Details matching screenshots
        payees = [
            {'name': 'BASA, AARON CHRISTIAN JUGAO', 'gross': 30782.40, 'tax': 3078.24}, # 10% tax
            {'name': 'CURA, SHERWIN RAFALLO', 'gross': 18014.40, 'tax': 900.72}, # 5% tax
            {'name': 'DELA CRUZ, JULIUS MAR LIWAG', 'gross': 29998.85, 'tax': 2999.88}, # 10% tax
            {'name': 'DEL MUNDO, JOEVEN EXEQUIEL VAZQUEZ', 'gross': 22214.40, 'tax': 2221.44}, # 10% tax
            {'name': 'SARDUAL, ELDEFONSO TANO JR.', 'gross': 22214.40, 'tax': 2221.44}, # 10% tax
            {'name': 'WADAWAD, RUBY ANN LUZ COPETE', 'gross': 22214.40, 'tax': 2221.44}, # 10% tax
        ]

        for p in payees:
            try:
                employee = Employee.objects.get(name=p['name'])
            except Employee.DoesNotExist:
                # Fallback if names differ slightly
                employee = Employee.objects.all().first()

            # Find matching active contract to link
            contract = Contract.objects.filter(employee=employee, is_active=True).first()

            # Dynamic particulars showing the dates based on screenshot
            particulars = f"Payment for services rendered for the period of JUNE 16-30, 2026 : {employee.name} (Period: 16Jun2026 - 30Jun2026)"
            
            DVPayeeDetail.objects.create(
                dv=dv,
                employee=employee,
                particulars=particulars,
                gross_amount=p['gross'],
                withholding_tax=p['tax'],
                net_amount=p['gross'] - p['tax'],
                contract=contract
            )

        messages.success(request, f"Voucher {dv.dv_number} and Request {obr.obr_number} successfully initiated from BULSA export!")
        return redirect('dv_view', pk=dv.pk)
    
    return render(request, 'payroll/create_transaction.html')

@login_required
def obr_view(request, pk):
    obr = get_object_or_404(ObligationRequest, pk=pk)
    # Get associated DV if any
    dv = obr.dvs.first()
    logs = obr.workflow_logs.all().order_by('changed_at')
    
    # Retrieve payee details from the linked DV for display
    payee_details = []
    if dv:
        payee_details = dv.payee_details.all()

    # Determine signatures to display based on status
    signatures = {
        'box_a': None, # Head, Requesting Unit (Isagani L. Bagus) - shows after UNIT_HEAD_APPROVED
        'box_b': None, # SBO Director (Noreen P. Escultura) - shows after APPROVED
    }

    if obr.status != ObRStatus.DRAFTED and obr.status != ObRStatus.SUBMITTED:
        # Unit head approved
        try:
            head_user = User.objects.get(username='unithead')
            signatures['box_a'] = head_user.profile.signature.url if head_user.profile.signature else None
        except User.DoesNotExist:
            pass

    if obr.status == ObRStatus.APPROVED:
        # SBO Approved
        try:
            sbo_dir = User.objects.get(username='sbo_director') # In screenshots Noreen P. Escultura is Director, SBO
            signatures['box_b'] = sbo_dir.profile.signature.url if sbo_dir.profile.signature else None
        except User.DoesNotExist:
            pass

    context = {
        'obr': obr,
        'dv': dv,
        'logs': logs,
        'payee_details': payee_details,
        'signatures': signatures,
        'ObRStatus': ObRStatus,
    }
    return render(request, 'payroll/obr_view.html', context)

@login_required
def dv_view(request, pk):
    dv = get_object_or_404(DisbursementVoucher, pk=pk)
    
    # Auto-repair/unblock if linked ObR is already approved
    if dv.status == DVStatus.WAITING_OBR and dv.obr and dv.obr.status == ObRStatus.APPROVED:
        dv.status = DVStatus.UNIT_HEAD_APPROVED
        dv.save()
        WorkflowLog.objects.create(
            dv=dv,
            status_from=DVStatus.WAITING_OBR,
            status_to=DVStatus.UNIT_HEAD_APPROVED,
            changed_by=request.user,
            remarks="ObR was approved: DV automatically unblocked upon inspection."
        )
        
    logs = dv.workflow_logs.all().order_by('changed_at')
    payee_details = dv.payee_details.all()
    
    # Check if pre-audit verification is complete (for pre_audit user)
    all_payees_verified = not payee_details.filter(is_contract_verified=False).exists()

    # Signatures configuration matching the boxes on the DV Form
    signatures = {
        'box_a': None, # Head, Unit (Isagani L. Bagus) - appears at UNIT_HEAD_APPROVED
        'box_c': None, # Director, SAO (Ronnie B. Pagal) - appears at SAO_DIRECTOR
        'box_d': None, # VP, OVPA (Augustus C. Resurreccion) - appears at APPROVED
    }

    # Show Box A signature if it has progressed past waiting obr
    if dv.status not in [DVStatus.DRAFTED, DVStatus.WAITING_OBR]:
        try:
            head = User.objects.get(username='unithead')
            signatures['box_a'] = head.profile.signature.url if head.profile.signature else None
        except User.DoesNotExist:
            pass

    # Show Box C signature (SAO Director)
    if dv.status in [DVStatus.SAO_DIRECTOR, DVStatus.APPROVED, DVStatus.RADA_DRAFT]:
        try:
            sao_dir = User.objects.get(username='director')
            signatures['box_c'] = sao_dir.profile.signature.url if sao_dir.profile.signature else None
        except User.DoesNotExist:
            pass

    # Show Box D signature (VP OVPA)
    if dv.status in [DVStatus.APPROVED, DVStatus.RADA_DRAFT]:
        try:
            vp = User.objects.get(username='vp')
            signatures['box_d'] = vp.profile.signature.url if vp.profile.signature else None
        except User.DoesNotExist:
            pass

    # Contract for split screen: load active payee based on request or default
    active_payee = None
    prev_payee_id = None
    next_payee_id = None
    
    if request.user.profile.role == Role.SAO_PREAUDIT:
        active_payee_id = request.GET.get('active_payee')
        if active_payee_id:
            try:
                active_payee = payee_details.get(pk=int(active_payee_id))
            except (ValueError, DVPayeeDetail.DoesNotExist):
                pass
        
        if not active_payee:
            active_payee = payee_details.filter(is_contract_verified=False).first()
        if not active_payee:
            active_payee = payee_details.first()
            
        # Calculate pagination
        if active_payee:
            payee_list = list(payee_details)
            try:
                curr_idx = payee_list.index(active_payee)
                if curr_idx > 0:
                    prev_payee_id = payee_list[curr_idx - 1].pk
                if curr_idx < len(payee_list) - 1:
                    next_payee_id = payee_list[curr_idx + 1].pk
            except ValueError:
                pass

    # Attach payroll details
    payee_details_with_payroll = []
    for detail in payee_details:
        payroll_item = None
        if dv.general_payroll:
            payroll_item = dv.general_payroll.items.filter(employee=detail.employee).first()
        payee_details_with_payroll.append({
            'detail': detail,
            'payroll_item': payroll_item
        })

    context = {
        'dv': dv,
        'logs': logs,
        'payee_details': payee_details,
        'payee_details_with_payroll': payee_details_with_payroll,
        'signatures': signatures,
        'all_payees_verified': all_payees_verified,
        'active_payee': active_payee,
        'prev_payee_id': prev_payee_id,
        'next_payee_id': next_payee_id,
        'DVStatus': DVStatus,
        'Role': Role,
    }
    return render(request, 'payroll/dv_view.html', context)

@login_required
def approve_obr(request, pk):
    obr = get_object_or_404(ObligationRequest, pk=pk)
    role = request.user.profile.role
    current_status = obr.status
    next_status = None
    remarks = request.POST.get('remarks', '')

    # Superuser admin override bypass
    if request.user.is_superuser:
        if current_status == ObRStatus.DRAFTED:
            next_status = ObRStatus.SUBMITTED
        elif current_status == ObRStatus.SUBMITTED:
            next_status = ObRStatus.UNIT_HEAD_APPROVED
        elif current_status == ObRStatus.UNIT_HEAD_APPROVED:
            next_status = ObRStatus.SBO_OFFICER_REVIEWED
        elif current_status == ObRStatus.SBO_OFFICER_REVIEWED:
            next_status = ObRStatus.SBO_SUPERVISOR_REVIEWED
        elif current_status == ObRStatus.SBO_SUPERVISOR_REVIEWED:
            next_status = ObRStatus.APPROVED
        remarks = remarks or "Superuser administrative override approval"
    else:
        # SBO workflow rules
        if current_status == ObRStatus.DRAFTED and role == Role.UNIT_AO:
            next_status = ObRStatus.SUBMITTED
            remarks = remarks or "Submitted ObR for endorsement"
        elif current_status == ObRStatus.SUBMITTED and role == Role.UNIT_HEAD:
            dv = obr.dvs.first()
            if not dv or not dv.general_payroll:
                messages.error(request, "Cannot approve: A certified General Payroll spreadsheet must be attached to this transaction before Unit Head approval.")
                return redirect('obr_view', pk=obr.pk)
            next_status = ObRStatus.UNIT_HEAD_APPROVED
            remarks = remarks or "Approved by Unit Head"
        elif current_status == ObRStatus.UNIT_HEAD_APPROVED and request.user.username == 'sbo_officer':
            next_status = ObRStatus.SBO_OFFICER_REVIEWED
            remarks = remarks or "Reviewed and certified by SBO Officer"
        elif current_status == ObRStatus.SBO_OFFICER_REVIEWED and request.user.username == 'sbo_supervisor':
            next_status = ObRStatus.SBO_SUPERVISOR_REVIEWED
            remarks = remarks or "Reviewed and certified by SBO Supervisor"
        elif current_status == ObRStatus.SBO_SUPERVISOR_REVIEWED and request.user.username == 'sbo_director':
            next_status = ObRStatus.APPROVED
            remarks = remarks or "Budget allocation APPROVED by SBO Director"

    if next_status:
        obr.status = next_status
        obr.save()

        # Log transition
        WorkflowLog.objects.create(
            obr=obr,
            status_from=current_status,
            status_to=next_status,
            changed_by=request.user,
            remarks=remarks
        )

        # UNBLOCK linked DV if ObR is APPROVED
        if next_status == ObRStatus.APPROVED:
            dv = obr.dvs.first()
            if dv and dv.status == DVStatus.WAITING_OBR:
                dv.status = DVStatus.UNIT_HEAD_APPROVED
                dv.save()
                WorkflowLog.objects.create(
                    dv=dv,
                    status_from=DVStatus.WAITING_OBR,
                    status_to=DVStatus.UNIT_HEAD_APPROVED,
                    changed_by=request.user,
                    remarks="ObR Approved: DV auto-released for SAO receiving."
                )
                messages.success(request, f"ObR approved! Linked Disbursement Voucher {dv.dv_number} has been automatically unblocked and sent to SAO.")

        messages.success(request, f"ObR status advanced to: {obr.get_status_display()}")
    else:
        messages.error(request, "You do not have permission to approve this document at its current stage.")

    return redirect('obr_view', pk=obr.pk)


@login_required
def approve_dv(request, pk):
    dv = get_object_or_404(DisbursementVoucher, pk=pk)
    role = request.user.profile.role
    current_status = dv.status
    next_status = None
    remarks = request.POST.get('remarks', '')

    # Superuser admin override bypass
    if request.user.is_superuser:
        if current_status == DVStatus.DRAFTED:
            if dv.obr and dv.obr.status == ObRStatus.APPROVED:
                next_status = DVStatus.UNIT_HEAD_APPROVED
            else:
                next_status = DVStatus.WAITING_OBR
        elif current_status == DVStatus.WAITING_OBR:
            next_status = DVStatus.UNIT_HEAD_APPROVED
        elif current_status == DVStatus.UNIT_HEAD_APPROVED:
            next_status = DVStatus.SAO_RECEIVED
        elif current_status == DVStatus.SAO_RECEIVED:
            next_status = DVStatus.SAO_PRE_AUDIT
        elif current_status == DVStatus.SAO_PRE_AUDIT:
            next_status = DVStatus.SAO_SUPERVISOR
        elif current_status == DVStatus.SAO_SUPERVISOR:
            next_status = DVStatus.CASH_MONITORED
        elif current_status == DVStatus.CASH_MONITORED:
            next_status = DVStatus.SAO_DIRECTOR
        elif current_status == DVStatus.SAO_DIRECTOR:
            next_status = DVStatus.APPROVED
        elif current_status == DVStatus.APPROVED:
            next_status = DVStatus.RADA_DRAFT
            dv.check_ada_no = request.POST.get('check_ada_no', 'ORS1-2026-07-00439')
            dv.payment_date = datetime.date.today()
            dv.payment_bank_account = 'UPS - COE Current Account (3072-1007-42)'
        remarks = remarks or "Superuser administrative override approval"
    else:
        # SAO workflow rules
        if current_status == DVStatus.DRAFTED and role == Role.UNIT_AO:
            if dv.obr and dv.obr.status == ObRStatus.APPROVED:
                next_status = DVStatus.UNIT_HEAD_APPROVED
                remarks = remarks or "DV Submitted: ObR already approved, auto-released to SAO"
            else:
                next_status = DVStatus.WAITING_OBR
                remarks = remarks or "DV Drafted: Awaiting associated ObR/BUR approval"
            
            # Auto-advance linked ObR to SUBMITTED if drafted
            if dv.obr and dv.obr.status == ObRStatus.DRAFTED:
                dv.obr.status = ObRStatus.SUBMITTED
                dv.obr.save()
                WorkflowLog.objects.create(
                    obr=dv.obr,
                    status_from=ObRStatus.DRAFTED,
                    status_to=ObRStatus.SUBMITTED,
                    changed_by=request.user,
                    remarks="ObR auto-submitted with DV"
                )
                
        elif current_status == DVStatus.UNIT_HEAD_APPROVED and role == Role.SAO_RECEIVING:
            next_status = DVStatus.SAO_RECEIVED
            remarks = remarks or "Voucher and digital attachments received by SAO"
        elif current_status == DVStatus.SAO_RECEIVED and role == Role.SAO_PREAUDIT:
            # Check if pre-audit verification of contracts is complete
            unverified_count = dv.payee_details.filter(is_contract_verified=False).count()
            if unverified_count > 0:
                messages.error(request, f"Cannot approve: {unverified_count} contracts are still unverified. Please verify and tag all payee contracts first.")
                return redirect('dv_view', pk=dv.pk)
            next_status = DVStatus.SAO_PRE_AUDIT
            remarks = remarks or "Pre-audit completed: digital contract attachments verified"
        elif current_status == DVStatus.SAO_PRE_AUDIT and role == Role.SAO_PAYROLLSUP:
            next_status = DVStatus.SAO_SUPERVISOR
            remarks = remarks or "Approved for Cash Monitoring"
        elif current_status == DVStatus.SAO_SUPERVISOR and role == Role.SAO_CASHMONI:
            next_status = DVStatus.CASH_MONITORED
            remarks = remarks or "Cash monitoring release finalized"
        elif current_status == DVStatus.CASH_MONITORED and role == Role.SAO_DIRECTOR:
            next_status = DVStatus.SAO_DIRECTOR
            remarks = remarks or "Approved by SAO Director"
        elif current_status == DVStatus.SAO_DIRECTOR and role == Role.OVPA_VP:
            next_status = DVStatus.APPROVED
            remarks = remarks or "Final approval granted by Vice President for Administration (OVPA)"
        elif current_status == DVStatus.APPROVED and role == Role.SCO_STAFF:
            next_status = DVStatus.RADA_DRAFT
            dv.check_ada_no = request.POST.get('check_ada_no', 'ORS1-2026-07-00439')
            dv.payment_date = datetime.date.today()
            dv.payment_bank_account = 'UPS - COE Current Account (3072-1007-42)'
            remarks = remarks or f"Draft RADA generated. Payment reference: {dv.check_ada_no}"


    if next_status:
        dv.status = next_status
        dv.save()

        # Log transition
        WorkflowLog.objects.create(
            dv=dv,
            status_from=current_status,
            status_to=next_status,
            changed_by=request.user,
            remarks=remarks
        )
        messages.success(request, f"DV status advanced to: {dv.get_status_display()}")
    else:
        messages.error(request, "You do not have permission to approve this document at its current stage.")

    return redirect('dv_view', pk=dv.pk)

@login_required
def verify_contract(request, pk):
    detail = get_object_or_404(DVPayeeDetail, pk=pk)
    # Check permissions
    if request.user.profile.role != Role.SAO_PREAUDIT:
        return HttpResponseForbidden("Only Pre-Audit Staff can verify contracts.")
    
    if request.method == 'POST':
        detail.is_contract_verified = True
        detail.verified_by = request.user
        detail.verified_at = timezone.now()
        detail.save()
        messages.success(request, f"Contract for {detail.employee.name} verified and tagged as digital attachment successfully!")
        
    return redirect('dv_view', pk=detail.dv.pk)

@login_required
def contracts_list(request):
    role = request.user.profile.role
    dept = request.user.profile.department
    
    # 1. Base Queryset
    if role in [Role.UNIT_AO, Role.UNIT_HEAD] and not request.user.is_superuser:
        contracts = Contract.objects.filter(employee__department=dept)
    else:
        contracts = Contract.objects.all()
        
    # 2. Search Query
    q = request.GET.get('q', '').strip()
    if q:
        contracts = contracts.filter(
            Q(employee__name__icontains=q) | 
            Q(designation__icontains=q)
        )
        
    # 3. Department Filter (for global/admin roles only)
    dept_filter = request.GET.get('dept', '').strip()
    if dept_filter and (role not in [Role.UNIT_AO, Role.UNIT_HEAD] or request.user.is_superuser):
        contracts = contracts.filter(employee__department=dept_filter)
        
    # 4. Engagement Type Filter
    eng_filter = request.GET.get('engagement', '').strip()
    if eng_filter:
        contracts = contracts.filter(rate_type=eng_filter)
        
    # 5. Year Filter
    year_filter = request.GET.get('year', '').strip()
    if year_filter:
        contracts = contracts.filter(start_date__year=int(year_filter))
        
    contracts = contracts.order_by('employee__name')
    
    # Pagination (25 items per page)
    paginator = Paginator(contracts, 25)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Build query string of active filters to preserve them in pagination links
    query_params = request.GET.copy()
    if 'page' in query_params:
        del query_params['page']
    url_params = query_params.urlencode()
    
    context = {
        'contracts': page_obj,
        'page_obj': page_obj,
        'url_params': url_params,
        'departments': Department.choices,
        'is_global': role not in [Role.UNIT_AO, Role.UNIT_HEAD] or request.user.is_superuser,
        'q': q,
        'selected_dept': dept_filter,
        'selected_eng': eng_filter,
        'selected_year': year_filter,
        'years': [2024, 2025, 2026, 2027],
    }
    return render(request, 'payroll/contracts_list.html', context)

@login_required
def employees_list(request):
    role = request.user.profile.role
    dept = request.user.profile.department
    
    # 1. Base Queryset
    if role in [Role.UNIT_AO, Role.UNIT_HEAD] and not request.user.is_superuser:
        employees = Employee.objects.filter(department=dept)
    else:
        employees = Employee.objects.all()

    # 2. Search Query
    q = request.GET.get('q', '').strip()
    if q:
        employees = employees.filter(
            Q(name__icontains=q) | 
            Q(tin__icontains=q) | 
            Q(email__icontains=q)
        )

    # 3. Department Filter (for global roles only)
    dept_filter = request.GET.get('dept', '').strip()
    if dept_filter and (role not in [Role.UNIT_AO, Role.UNIT_HEAD] or request.user.is_superuser):
        employees = employees.filter(department=dept_filter)

    # 4. Engagement Type Filter
    eng_filter = request.GET.get('engagement', '').strip()
    if eng_filter:
        employees = employees.filter(contracts__rate_type=eng_filter, contracts__is_active=True)

    employees = employees.order_by('name')
    
    # Pagination (25 items per page)
    paginator = Paginator(employees, 25)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    # Build query string of active filters to preserve them in pagination links
    query_params = request.GET.copy()
    if 'page' in query_params:
        del query_params['page']
    url_params = query_params.urlencode()
    
    context = {
        'employees': page_obj,
        'page_obj': page_obj,
        'url_params': url_params,
        'departments': Department.choices,
        'is_global': role not in [Role.UNIT_AO, Role.UNIT_HEAD] or request.user.is_superuser,
        'q': q,
        'selected_dept': dept_filter,
        'selected_eng': eng_filter,
    }
    return render(request, 'payroll/employees_list.html', context)

@login_required
def general_payroll_list(request):
    payrolls = GeneralPayroll.objects.all().order_by('-created_at')
    return render(request, 'payroll/general_payroll_list.html', {'payrolls': payrolls})

@login_required
def general_payroll_create(request):
    profile = request.user.profile
    
    # Restrict to Unit AO per department
    if profile.role != Role.UNIT_AO and not request.user.is_superuser:
        return HttpResponseForbidden("Only Unit Administrative Officers can create general payroll sheets.")

    dept = profile.department
    employees = Employee.objects.filter(department=dept)

    # Gather active contracts for preview
    emp_contracts = []
    for emp in employees:
        contract = Contract.objects.filter(employee=emp, is_active=True).first()
        if contract:
            emp_contracts.append({
                'employee': emp,
                'contract': contract,
                'salary_basis': contract.monthly_rate if contract.rate_type == ContractType.MONTHLY else contract.daily_rate * 22
            })

    if request.method == 'POST':
        payroll_name = request.POST.get('payroll_name', 'General Payroll')
        fund = request.POST.get('fund', '101')
        bank = request.POST.get('bank', 'LBP')
        period_start = request.POST.get('period_start')
        period_end = request.POST.get('period_end')

        if not period_start or not period_end:
            messages.error(request, "Please enter validity start and end dates.")
            return redirect('general_payroll_create')

        # Create main payroll registry
        payroll = GeneralPayroll.objects.create(
            payroll_name=payroll_name,
            department=dept,
            fund=fund,
            bank=bank,
            period_start=period_start,
            period_end=period_end,
            created_by=request.user
        )

        # Loop through employees and grab inputs
        included_ids = request.POST.getlist('included_employees')
        for emp_id_str in included_ids:
            emp = get_object_or_404(Employee, id=int(emp_id_str))
            salary_basis = request.POST.get(f'salary_basis_{emp.pk}', 0.00)
            no_of_days = request.POST.get(f'no_of_days_{emp.pk}', 11.00)
            gross_salary = request.POST.get(f'gross_salary_{emp.pk}', 0.00)
            philhealth = request.POST.get(f'philhealth_{emp.pk}', 0.00)
            sss = request.POST.get(f'sss_{emp.pk}', 0.00)
            nvat = request.POST.get(f'nvat_{emp.pk}', 0.00)
            ewt_5 = request.POST.get(f'ewt_5_{emp.pk}', 0.00)
            ewt_10 = request.POST.get(f'ewt_10_{emp.pk}', 0.00)
            total_deductions = request.POST.get(f'total_deductions_{emp.pk}', 0.00)
            net_salary = request.POST.get(f'net_salary_{emp.pk}', 0.00)

            GeneralPayrollItem.objects.create(
                general_payroll=payroll,
                employee=emp,
                salary_basis=salary_basis,
                no_of_days=no_of_days,
                gross_salary=gross_salary,
                philhealth_deductions=philhealth,
                sss=sss,
                nvat_3=nvat,
                ewt_5=ewt_5,
                ewt_10=ewt_10,
                total_deductions=total_deductions,
                net_salary=net_salary
            )

        messages.success(request, f"General Payroll for {payroll.get_department_display()} successfully generated!")
        return redirect('general_payroll_detail', pk=payroll.pk)

    import json
    all_employees = Employee.objects.all().order_by('name')
    all_contracts_data = []
    for emp in all_employees:
        contract = Contract.objects.filter(employee=emp, is_active=True).first()
        salary_basis = float(contract.monthly_rate if contract.rate_type == ContractType.MONTHLY else contract.daily_rate * 22) if contract else 0.00
        all_contracts_data.append({
            'id': emp.id,
            'name': emp.name,
            'tin': emp.tin,
            'bank_name': emp.bank_name,
            'account_number': emp.account_number,
            'salary_basis': salary_basis,
            'designation': contract.designation if contract else "COS Staff",
            'department': emp.get_department_display()
        })
    all_contracts_json = json.dumps(all_contracts_data)

    context = {
        'department_display': profile.get_department_display(),
        'emp_contracts': emp_contracts,
        'all_employees': all_employees,
        'all_contracts_json': all_contracts_json,
    }
    return render(request, 'payroll/general_payroll_create.html', context)

@login_required
def general_payroll_detail(request, pk):
    payroll = get_object_or_404(GeneralPayroll, pk=pk)
    
    # Handle POST Actions (Add/Remove)
    if request.method == 'POST':
        profile = request.user.profile
        if profile.role not in [Role.UNIT_AO, Role.UNIT_HEAD] and not request.user.is_superuser:
            return HttpResponseForbidden("Only authorized department officers can modify the payroll sheet.")
            
        action = request.POST.get('action')
        if action == 'add_employee':
            emp_id = request.POST.get('employee_id')
            emp = get_object_or_404(Employee, id=emp_id)
            
            # Check if already exists in this payroll
            if payroll.items.filter(employee=emp).exists():
                messages.error(request, f"{emp.name} is already in this payroll sheet.")
            else:
                # Pre-calculate defaults
                contract = Contract.objects.filter(employee=emp, is_active=True).first()
                if contract:
                    salary_basis = float(contract.monthly_rate if contract.rate_type == ContractType.MONTHLY else contract.daily_rate * 22)
                else:
                    salary_basis = 0.00
                
                no_of_days = 11.00
                gross_salary = (salary_basis / 22) * no_of_days
                
                # Default deductions
                philhealth = gross_salary * 0.05
                ewt_5 = gross_salary * 0.05
                total_deductions = philhealth + ewt_5
                net_salary = gross_salary - total_deductions
                
                GeneralPayrollItem.objects.create(
                    general_payroll=payroll,
                    employee=emp,
                    salary_basis=salary_basis,
                    no_of_days=no_of_days,
                    gross_salary=gross_salary,
                    philhealth_deductions=philhealth,
                    ewt_5=ewt_5,
                    total_deductions=total_deductions,
                    net_salary=net_salary
                )
                messages.success(request, f"Successfully added {emp.name} to the payroll sheet.")
            return redirect('general_payroll_detail', pk=pk)
            
        elif action == 'delete_employee':
            item_id = request.POST.get('item_id')
            item = get_object_or_404(GeneralPayrollItem, id=item_id, general_payroll=payroll)
            emp_name = item.employee.name
            item.delete()
            messages.success(request, f"Removed {emp_name} from the payroll sheet.")
            return redirect('general_payroll_detail', pk=pk)

    items = payroll.items.all().order_by('employee__name')
    
    # Filter out employees already in this payroll
    existing_emp_ids = items.values_list('employee_id', flat=True)
    available_employees = Employee.objects.all().exclude(id__in=existing_emp_ids).order_by('name')
    
    # Calculate totals
    totals = {
        'gross': sum(i.gross_salary for i in items),
        'philhealth': sum(i.philhealth_deductions for i in items),
        'sss': sum(i.sss for i in items),
        'nvat': sum(i.nvat_3 for i in items),
        'ewt_5': sum(i.ewt_5 for i in items),
        'ewt_10': sum(i.ewt_10 for i in items),
        'deductions': sum(i.total_deductions for i in items),
        'net': sum(i.net_salary for i in items),
    }

    # Find the acting chief signature
    signatures = {
        'head': None
    }
    try:
        head = User.objects.get(username='unithead')
        signatures['head'] = head.profile.signature.url if head.profile.signature else None
    except User.DoesNotExist:
        pass

    context = {
        'payroll': payroll,
        'items': items,
        'totals': totals,
        'signatures': signatures,
        'available_employees': available_employees,
    }
    return render(request, 'payroll/general_payroll_detail.html', context)


@login_required
def employee_create(request):
    profile = request.user.profile
    if profile.role not in [Role.UNIT_AO, Role.UNIT_HEAD] and not request.user.is_superuser:
        return HttpResponseForbidden("Only department officers can register personnel.")
        
    if request.method == 'POST':
        name = request.POST.get('name')
        tin = request.POST.get('tin')
        bank_name = request.POST.get('bank_name', 'LBP')
        account_number = request.POST.get('account_number')
        unit = request.POST.get('unit', '')
        address = request.POST.get('address', '')
        govt_id = request.POST.get('govt_id', '')
        govt_id_details = request.POST.get('govt_id_details', '')
        
        if request.user.is_superuser:
            department = request.POST.get('department')
        else:
            department = profile.department
            
        employee = Employee.objects.create(
            name=name,
            tin=tin,
            bank_name=bank_name,
            account_number=account_number,
            department=department,
            unit=unit,
            address=address,
            govt_id=govt_id,
            govt_id_details=govt_id_details
        )
        
        designation = request.POST.get('designation', 'COS Staff')
        rate_type = request.POST.get('rate_type', ContractType.MONTHLY)
        monthly_rate = float(request.POST.get('monthly_rate') or 0.00)
        daily_rate = float(request.POST.get('daily_rate') or 0.00)
        start_date = request.POST.get('start_date') or '2026-01-01'
        end_date = request.POST.get('end_date') or '2026-06-30'
        
        duties_list = [d.strip() for d in request.POST.getlist('duties[]') if d.strip()]
        duties = "\n".join(duties_list) if duties_list else 'Perform general administrative functions.'
        funding_source = request.POST.get('funding_source', 'UPS SPMO MOOE')
        head_of_unit_name = request.POST.get('head_of_unit_name', 'ISAGANI L. BAGUS')
        head_of_unit_position = request.POST.get('head_of_unit_position', 'Acting Chief SSPMO')
        head_of_unit_id = request.POST.get('head_of_unit_id', 'NO4-03-000009')
        head_of_unit_id_details = request.POST.get('head_of_unit_id_details', 'June 8, 2022 / Quezon City')
        witness1_name = request.POST.get('witness1_name', 'MARK JOSHUA M. PEDROSA')
        witness1_position = request.POST.get('witness1_position', 'Administrative Assistant I')
        witness2_name = request.POST.get('witness2_name', 'JULIUS MAR L. DELA CRUZ')
        witness2_position = request.POST.get('witness2_position', 'Junior Office Manager')

        Contract.objects.create(
            employee=employee,
            designation=designation,
            rate_type=rate_type,
            monthly_rate=monthly_rate,
            daily_rate=daily_rate,
            start_date=start_date,
            end_date=end_date,
            is_active=True,
            duties=duties,
            funding_source=funding_source,
            head_of_unit_name=head_of_unit_name,
            head_of_unit_position=head_of_unit_position,
            head_of_unit_id=head_of_unit_id,
            head_of_unit_id_details=head_of_unit_id_details,
            witness1_name=witness1_name,
            witness1_position=witness1_position,
            witness2_name=witness2_name,
            witness2_position=witness2_position
        )
        
        messages.success(request, f"Employee {employee.name} registered successfully!")
        return redirect('employees_list')
        
    context = {
        'departments': Department.choices,
        'is_superuser': request.user.is_superuser,
        'profile_department': profile.get_department_display(),
        'ContractType': ContractType,
    }
    return render(request, 'payroll/employee_form.html', context)

@login_required
def employee_edit(request, pk):
    employee = get_object_or_404(Employee, pk=pk)
    profile = request.user.profile
    
    if employee.department != profile.department and not request.user.is_superuser:
        return HttpResponseForbidden("You do not have permission to edit this employee because they belong to another department.")
        
    contract = employee.contracts.filter(is_active=True).first()
    
    if request.method == 'POST':
        employee.name = request.POST.get('name')
        employee.tin = request.POST.get('tin')
        employee.bank_name = request.POST.get('bank_name')
        employee.account_number = request.POST.get('account_number')
        employee.unit = request.POST.get('unit')
        employee.address = request.POST.get('address', '')
        employee.govt_id = request.POST.get('govt_id', '')
        employee.govt_id_details = request.POST.get('govt_id_details', '')
        
        if request.user.is_superuser:
            employee.department = request.POST.get('department')
            
        employee.save()
        
        designation = request.POST.get('designation', 'COS Staff')
        rate_type = request.POST.get('rate_type', ContractType.MONTHLY)
        monthly_rate = float(request.POST.get('monthly_rate') or 0.00)
        daily_rate = float(request.POST.get('daily_rate') or 0.00)
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')
        
        duties_list = [d.strip() for d in request.POST.getlist('duties[]') if d.strip()]
        duties = "\n".join(duties_list) if duties_list else 'Perform general administrative functions.'
        funding_source = request.POST.get('funding_source', 'UPS SPMO MOOE')
        head_of_unit_name = request.POST.get('head_of_unit_name', 'ISAGANI L. BAGUS')
        head_of_unit_position = request.POST.get('head_of_unit_position', 'Acting Chief SSPMO')
        head_of_unit_id = request.POST.get('head_of_unit_id', 'NO4-03-000009')
        head_of_unit_id_details = request.POST.get('head_of_unit_id_details', 'June 8, 2022 / Quezon City')
        witness1_name = request.POST.get('witness1_name', 'MARK JOSHUA M. PEDROSA')
        witness1_position = request.POST.get('witness1_position', 'Administrative Assistant I')
        witness2_name = request.POST.get('witness2_name', 'JULIUS MAR L. DELA CRUZ')
        witness2_position = request.POST.get('witness2_position', 'Junior Office Manager')

        if contract:
            contract.designation = designation
            contract.rate_type = rate_type
            contract.monthly_rate = monthly_rate
            contract.daily_rate = daily_rate
            if start_date:
                contract.start_date = start_date
            if end_date:
                contract.end_date = end_date
            
            contract.duties = duties
            contract.funding_source = funding_source
            contract.head_of_unit_name = head_of_unit_name
            contract.head_of_unit_position = head_of_unit_position
            contract.head_of_unit_id = head_of_unit_id
            contract.head_of_unit_id_details = head_of_unit_id_details
            contract.witness1_name = witness1_name
            contract.witness1_position = witness1_position
            contract.witness2_name = witness2_name
            contract.witness2_position = witness2_position
            contract.save()
        else:
            Contract.objects.create(
                employee=employee,
                designation=designation,
                rate_type=rate_type,
                monthly_rate=monthly_rate,
                daily_rate=daily_rate,
                start_date=start_date or '2026-01-01',
                end_date=end_date or '2026-06-30',
                is_active=True,
                duties=duties,
                funding_source=funding_source,
                head_of_unit_name=head_of_unit_name,
                head_of_unit_position=head_of_unit_position,
                head_of_unit_id=head_of_unit_id,
                head_of_unit_id_details=head_of_unit_id_details,
                witness1_name=witness1_name,
                witness1_position=witness1_position,
                witness2_name=witness2_name,
                witness2_position=witness2_position
            )
            
        messages.success(request, f"Employee {employee.name} updated successfully!")
        return redirect('employees_list')
        
    context = {
        'employee': employee,
        'contract': contract,
        'contract_duties_list': [d.strip() for d in contract.duties.split('\n') if d.strip()] if contract and contract.duties else [],
        'departments': Department.choices,
        'is_superuser': request.user.is_superuser,
        'profile_department': profile.get_department_display(),
        'ContractType': ContractType,
    }
    return render(request, 'payroll/employee_form.html', context)

@login_required
def employee_delete(request, pk):
    if not request.user.is_superuser:
        return HttpResponseForbidden("Only superusers can delete employee records.")
        
    employee = get_object_or_404(Employee, pk=pk)
    if request.method == 'POST':
        name = employee.name
        employee.delete()
        messages.success(request, f"Employee {name} deleted successfully.")
        return redirect('employees_list')
        
    return render(request, 'payroll/employee_confirm_delete.html', {'employee': employee})

@login_required
def memos_viewer(request):
    return render(request, 'payroll/memos_viewer.html')

@login_required
def generate_contract_docx(request, employee_id):
    employee = get_object_or_404(Employee, pk=employee_id)
    profile = request.user.profile
    
    if employee.department != profile.department and not request.user.is_superuser:
        return HttpResponseForbidden("You do not have permission to generate this contract.")
        
    contract = employee.contracts.filter(is_active=True).first()
    if not contract:
        messages.error(request, "This employee does not have an active contract.")
        return redirect('contracts_list')

    if not contract.is_viewable:
        messages.error(request, "Contract document generation is only available for engagements ending on or after July 1, 2026.")
        return redirect('contracts_list')

    if contract.rate_type == ContractType.DAILY:
        template_path = os.path.join(settings.MEDIA_ROOT, 'templates', 'daily-template.docx')
    else:
        template_path = os.path.join(settings.MEDIA_ROOT, 'templates', 'monthly-template.docx')

    if not os.path.exists(template_path):
        messages.error(request, f"Contract template file missing at media root: {template_path}")
        return redirect('contracts_list')

    try:
        doc = docx.Document(template_path)
    except Exception as e:
        messages.error(request, f"Error opening template: {str(e)}")
        return redirect('contracts_list')

        return redirect('contracts_list')

    # Word numbers helper
    def num_to_words(number):
        units = ["", "one", "two", "three", "four", "five", "six", "seven", "eight", "nine", "ten", "eleven", "twelve", "thirteen", "fourteen", "fifteen", "sixteen", "seventeen", "eighteen", "nineteen"]
        tens = ["", "", "twenty", "thirty", "forty", "fifty", "sixty", "seventy", "eighty", "ninety"]
        
        def helper(n):
            if n < 20:
                return units[int(n)]
            elif n < 100:
                return tens[int(n // 10)] + ("-" + units[int(n % 10)] if n % 10 != 0 else "")
            elif n < 1000:
                return units[int(n // 100)] + " hundred" + (" and " + helper(n % 100) if n % 100 != 0 else "")
            elif n < 1000000:
                return helper(n // 1000) + " thousand " + (helper(n % 1000) if n % 1000 != 0 else "")
            return str(n)
            
        num_val = float(number)
        num_int = int(num_val)
        num_dec = int(round((num_val - num_int) * 100))
        
        words = helper(num_int)
        if num_dec > 0:
            words += f" pesos and {num_dec}/100"
        else:
            words += " pesos"
        return words.capitalize()

    # Replacements setup
    replacements = {
        "[HEAD OF UNIT]": contract.head_of_unit_name.upper(),
        "ISAGANI L. BAGUS": contract.head_of_unit_name.upper(),
        "Acting Chief SSPMO": contract.head_of_unit_position,
        "Acting Chief, UP System SPMO": contract.head_of_unit_position,
        "[Position]": contract.designation,
        "[Office/Unit]": employee.get_department_display(),
        "[FULL NAME of CONTRACT of SERVICE PERSONNEL]": employee.name.upper(),
        "[Name of SECOND PARTY]": employee.name.upper(),
        "[Name]": employee.name.upper(),
        "AARON CHRISTIAN J. BASA": employee.name.upper(),
        "[Complete Address]": employee.address or "Diliman, Quezon City, Metro Manila",
        "BLK 61 LOT 44, LAKAMBINI ST., LAGRO QC": employee.address or "Diliman, Quezon City, Metro Manila",
        "[Start Date]": contract.start_date.strftime("%d %B %Y").upper() if contract.start_date else "",
        "[End Date]": contract.end_date.strftime("%d %B %Y").upper() if contract.end_date else "",
        "01 JULY 2026": contract.start_date.strftime("%d %B %Y").upper() if contract.start_date else "",
        "31 DECEMBER 2026": contract.end_date.strftime("%d %B %Y").upper() if contract.end_date else "",
        "[Funding Source]": contract.funding_source,
        "UPS SPMO MOOE": contract.funding_source,
        "MARK JOSHUA M. PEDROSA": contract.witness1_name,
        "Administrative Assistant I": contract.witness1_position,
        "JULIUS MAR L. DELA CRUZ": contract.witness2_name,
        "Junior Office Manager": contract.witness2_position,
        "NO4-03-000009": contract.head_of_unit_id,
        "June 8,2022 /Quezon City": contract.head_of_unit_id_details,
        "June 8, 2022 / Quezon City": contract.head_of_unit_id_details,
        "KO1-14-000493": employee.govt_id,
        "January 6, 2025/Cabadbaran City": employee.govt_id_details,
    }

    if contract.rate_type == ContractType.DAILY:
        daily_val = float(contract.daily_rate)
        prem_val = daily_val * 0.20
        total_val = daily_val + prem_val
        
        replacements.update({
            "[Daily Rate in Words]": num_to_words(daily_val),
            "Php X,XXX.XX": f"Php {daily_val:,.2f}",
            "Php XXX.XX": f"Php {prem_val:,.2f}",
            "[Amount in Words of Daily Premium]": num_to_words(prem_val),
            "[Total Daily Rate in Words]": num_to_words(total_val),
        })
    else:
        monthly_val = float(contract.monthly_rate)
        prem_val = monthly_val * 0.20
        total_val = monthly_val + prem_val
        
        replacements.update({
            "[Amount in Words]": num_to_words(monthly_val),
            "fifty thousand eight hundred eighteen pesos": num_to_words(monthly_val).lower(),
            "fifty-three thousand eight hundred eighteen pesos": num_to_words(monthly_val).lower(),
            "P53,818.00": f"P{monthly_val:,.2f}",
            "Php XX,XXX.XX": f"Php {monthly_val:,.2f}",
            "[Amount in Words of Premium]": num_to_words(prem_val),
            "Ten Thousand Seven Hundred Sixty-Three Pesos": num_to_words(prem_val),
            "P10,763.00": f"P{prem_val:,.2f}",
            "Php X,XXX.XX": f"Php {prem_val:,.2f}",
            "[Total Amount in Words]": num_to_words(total_val),
        })

    # Paragraph replace
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                for run in p.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, val)
                if key in p.text:
                    p.text = p.text.replace(key, val)

    # Table cell replace
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in replacements.items():
                        if key in p.text:
                            for run in p.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, val)
                            if key in p.text:
                                p.text = p.text.replace(key, val)

    # Replace duties (bullets list)
    duties_list = [d.strip() for d in contract.duties.split('\n') if d.strip()]
    if duties_list:
        duties_idx = 0
        for p in doc.paragraphs:
            if "[Function / Service" in p.text or "Draft funding and budget clearance requests" in p.text:
                if duties_idx < len(duties_list):
                    p.text = duties_list[duties_idx]
                    duties_idx += 1
                else:
                    p.text = "" # Clear unused bullet placeholders

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    filename = f"Contract_{employee.name.replace(' ', '_')}_{contract.get_rate_type_display().replace(' ', '_')}.docx"
    response = FileResponse(
        file_stream, 
        as_attachment=True, 
        filename=filename,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    return response


@login_required
def admin_command_center(request):
    if not request.user.is_superuser:
        return HttpResponseForbidden("You do not have permission to access the Command Center.")

    import csv
    import platform
    import sys
    from django.http import HttpResponse
    from django.db import connection

    # Handle POST Actions
    if request.method == 'POST':
        action = request.POST.get('action')
        
        if action == 'update_user_role':
            user_id = request.POST.get('user_id')
            new_role = request.POST.get('role')
            new_dept = request.POST.get('department')
            target_user = get_object_or_404(User, id=user_id)
            profile = target_user.profile
            if new_role:
                profile.role = new_role
            if new_dept:
                profile.department = new_dept
            profile.save()
            messages.success(request, f"Successfully updated profile for user {target_user.username}.")
            return redirect('admin_command_center')
            
        elif action == 'override_status':
            tx_type = request.POST.get('tx_type')
            tx_id = request.POST.get('tx_id')
            new_status = request.POST.get('new_status')
            remarks = request.POST.get('remarks', 'Status overridden via Command Center.')
            
            if tx_type == 'OBR':
                obr = get_object_or_404(ObligationRequest, id=tx_id)
                old_status = obr.status
                obr.status = new_status
                obr.save()
                
                WorkflowLog.objects.create(
                    obr=obr,
                    status_from=old_status,
                    status_to=new_status,
                    changed_by=request.user,
                    remarks=remarks
                )
                messages.success(request, f"Obligation Request {obr.obr_number} status overridden to {new_status}.")
            elif tx_type == 'DV':
                dv = get_object_or_404(DisbursementVoucher, id=tx_id)
                old_status = dv.status
                dv.status = new_status
                dv.save()
                
                WorkflowLog.objects.create(
                    dv=dv,
                    status_from=old_status,
                    status_to=new_status,
                    changed_by=request.user,
                    remarks=remarks
                )
                messages.success(request, f"Disbursement Voucher {dv.dv_number} status overridden to {new_status}.")
            return redirect('admin_command_center')
            
        elif action == 'export_report':
            report_type = request.POST.get('report_type')
            response = HttpResponse(content_type='text/csv')
            
            if report_type == 'OBR':
                response['Content-Disposition'] = 'attachment; filename="obr_summary_report.csv"'
                writer = csv.writer(response)
                writer.writerow(['ID', 'ObR Number', 'Date', 'Requesting Unit', 'Expense Class', 'Total Amount', 'Status'])
                for obr in ObligationRequest.objects.all().order_by('-transaction_date'):
                    writer.writerow([obr.id, obr.obr_number, obr.transaction_date, obr.requesting_unit, obr.expense_class, obr.total_amount, obr.get_status_display()])
            else:
                response['Content-Disposition'] = 'attachment; filename="dv_summary_report.csv"'
                writer = csv.writer(response)
                writer.writerow(['ID', 'DV Number', 'Date', 'Requesting Unit', 'Mode of Payment', 'Total Amount', 'Status'])
                for dv in DisbursementVoucher.objects.all().order_by('-transaction_date'):
                    writer.writerow([dv.id, dv.dv_number, dv.transaction_date, dv.requesting_unit, dv.mode_of_payment, dv.total_amount, dv.get_status_display()])
            return response

    # GET requests - Collect status and metrics
    # System Stability
    sys_info = {
        'os': platform.system() + ' ' + platform.release(),
        'python_version': sys.version.split(' ')[0],
        'django_version': '6.0.6',
        'db_backend': connection.vendor,
    }
    
    # DB Size and connection check
    db_size = "Unknown"
    db_healthy = False
    try:
        if connection.vendor == 'postgresql':
            with connection.cursor() as cursor:
                cursor.execute("SELECT pg_size_pretty(pg_database_size(current_database()));")
                db_size = cursor.fetchone()[0]
                db_healthy = True
        elif connection.vendor == 'sqlite':
            db_path = settings.DATABASES['default']['NAME']
            if os.path.exists(db_path):
                size_bytes = os.path.getsize(db_path)
                db_size = f"{size_bytes / (1024*1024):.2f} MB"
                db_healthy = True
    except Exception as e:
        db_size = f"Error reading DB: {str(e)}"
        
    # Backend Storage Status
    media_path = settings.MEDIA_ROOT
    media_size = "0 MB"
    media_count = 0
    try:
        if os.path.exists(media_path):
            total_size = 0
            for dirpath, dirnames, filenames in os.walk(media_path):
                for f in filenames:
                    fp = os.path.join(dirpath, f)
                    if not os.path.islink(fp):
                        total_size += os.path.getsize(fp)
                        media_count += 1
            media_size = f"{total_size / (1024*1024):.2f} MB"
    except Exception:
        pass

    supabase_configured = bool(settings.SUPABASE_URL and settings.SUPABASE_KEY)
    
    storage_info = {
        'supabase_configured': supabase_configured,
        'supabase_url': settings.SUPABASE_URL if supabase_configured else "Not Set",
        'media_count': media_count,
        'media_size': media_size,
    }

    # Workflow / Approval Flow Statistics
    obr_status_stats = {}
    for code, label in ObRStatus.choices:
        obr_status_stats[label] = {
            'code': code,
            'count': ObligationRequest.objects.filter(status=code).count()
        }
        
    dv_status_stats = {}
    for code, label in DVStatus.choices:
        dv_status_stats[label] = {
            'code': code,
            'count': DisbursementVoucher.objects.filter(status=code).count()
        }
        
    recent_logs = WorkflowLog.objects.all().order_by('-changed_at')[:15]

    # User Management
    all_users = User.objects.all().select_related('profile').order_by('username')

    # General Transaction History Search and Filter
    search_query = request.GET.get('q', '')
    filter_type = request.GET.get('type', '')
    filter_status = request.GET.get('status', '')

    tx_list = []
    
    # Query Obligation Requests
    obrs = ObligationRequest.objects.all()
    if search_query:
        obrs = obrs.filter(Q(obr_number__icontains=search_query) | Q(requesting_unit__icontains=search_query))
    if filter_status:
        obrs = obrs.filter(status=filter_status)
    if filter_type == '' or filter_type == 'OBR':
        for o in obrs:
            tx_list.append({
                'type': 'OBR',
                'id': o.id,
                'number': o.obr_number,
                'date': o.transaction_date,
                'requesting_unit': o.requesting_unit,
                'amount': o.total_amount,
                'status': o.status,
                'status_display': o.get_status_display(),
                'url': f"/obr/{o.id}/",
                'status_choices': ObRStatus.choices
            })

    # Query Disbursement Vouchers
    dvs = DisbursementVoucher.objects.all()
    if search_query:
        dvs = dvs.filter(Q(dv_number__icontains=search_query) | Q(requesting_unit__icontains=search_query))
    if filter_status:
        dvs = dvs.filter(status=filter_status)
    if filter_type == '' or filter_type == 'DV':
        for d in dvs:
            tx_list.append({
                'type': 'DV',
                'id': d.id,
                'number': d.dv_number,
                'date': d.transaction_date,
                'requesting_unit': d.requesting_unit,
                'amount': d.total_amount,
                'status': d.status,
                'status_display': d.get_status_display(),
                'url': f"/dv/{d.id}/",
                'status_choices': DVStatus.choices
            })

    # Sort combined transaction history by date desc
    tx_list.sort(key=lambda x: x['date'], reverse=True)

    # Paginate results
    paginator = Paginator(tx_list, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    context = {
        'sys_info': sys_info,
        'db_healthy': db_healthy,
        'db_size': db_size,
        'storage_info': storage_info,
        'obr_status_stats': obr_status_stats,
        'dv_status_stats': dv_status_stats,
        'recent_logs': recent_logs,
        'all_users': all_users,
        'page_obj': page_obj,
        'search_query': search_query,
        'filter_type': filter_type,
        'filter_status': filter_status,
        'role_choices': Role.choices,
        'dept_choices': Department.choices,
        'obr_statuses': ObRStatus.choices,
        'dv_statuses': DVStatus.choices,
    }

    return render(request, 'payroll/admin_command_center.html', context)



